# ==================================
# Это ужас и кошммар, я хз как сделать нормальное оформление docx
# ==================================

import io
import uuid
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
 
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
 
from app.models.sessions import Session
from app.models.test import Test
from app.models.answers import Answer
from app.models.question import Question, QuestionType
from app.models.formula import Formula
from app.models.sections import Section
 
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# =====================================
# ПОЛУЧЕНИЕ ДАННЫХ СЕССИИ
# =====================================


async def load_session_data(db: AsyncSession, session_id: uuid.UUID) -> dict:
    """
    загружает данные для генерации отчёта
    сессию, тест, ответы, вопросы, метрики.
    """
    # сессия + тест + формулы
    session_result = await db.execute(
        select(Session)
        .where(Session.id == session_id)
        .options(
            selectinload(Session.test).selectinload(Test.formula),
            selectinload(Session.test).selectinload(Test.sections).selectinload(Section.questions),
            selectinload(Session.answers).selectinload(Answer.question),
        )
    )
    session: Session | None = session_result.scalar_one_or_none()
    if not session:
        raise ValueError(f"Сессия {session_id} не найдена")
 
    # индексы для быстрого доступа
    answers_by_question: dict[uuid.UUID, Answer] = {
        a.question_id: a for a in session.answers
    }
 
    questions_by_id: dict[uuid.UUID, Question] = {
        q.id: q
        for section in session.test.sections
        for q in section.questions
    }
 
    return {
        "session":             session,
        "test":                session.test,
        "answers_by_question": answers_by_question,
        "questions_by_id":     questions_by_id,
        "metrics":             session.metrics or {},
    }
 
 
# =====================================
# ГЕНЕРАЦИЯ ГРАФИКОВ
# =====================================
 
def make_bar_chart(labels: list[str], values: list[float], title: str) -> io.BytesIO:
    """рисует столбчатый график через и сохраняет его в .png"""
    fig, ax = plt.subplots(figsize=(6, 3))
    x = np.arange(len(labels))
    ax.bar(x, values, color="#4472C4")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_title(title, fontsize=11)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf
 
 
def make_radar_chart(labels: list[str], values: list[float], title: str) -> io.BytesIO:
    n = len(labels)
    if n < 3:
        # радар с меньше чем 3 осями не имеет смысла поэтому делаем бар
        return make_bar_chart(labels, values, title)
 
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    values_plot = values + [values[0]]
    angles += [angles[0]]
 
    fig, ax = plt.subplots(figsize=(4, 4), subplot_kw={"polar": True})
    ax.plot(angles, values_plot, color="#4472C4", linewidth=2)
    ax.fill(angles, values_plot, color="#4472C4", alpha=0.2)
    ax.set_thetagrids(np.degrees(angles[:-1]), labels, fontsize=9) # type: ignore
    ax.set_title(title, fontsize=11, pad=15)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf
 
 
def make_chart(chart_type: str, labels: list[str], values: list[float], title: str) -> io.BytesIO:
    if chart_type == "radar":
        return make_radar_chart(labels, values, title)
    return make_bar_chart(labels, values, title)
 
 
# =====================================
# РЕНДЕР БЛОКОВ
# =====================================
 
def render_text_block(doc: DocxDocument, block: dict) -> None:
    """просто текст"""
    content = block.get("content", "")
    if content:
        doc.add_paragraph(content)
 
 
def render_metric_block(doc: DocxDocument, block: dict, metrics: dict) -> None:
    """значение метрики + что она значит"""
    metric_id = block.get("metric_id", "")
    metric_data = metrics.get(metric_id)
    if not metric_data:
        return
 
    name = metric_data.get("name", metric_id)
    value = metric_data.get("value")
    interpretation = metric_data.get("interpretation", "")
 
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(str(value))
    if interpretation:
        p.add_run(f"  — {interpretation}").italic = True
 
 
def render_chart_block(doc: DocxDocument, block: dict, metrics: dict) -> None:
    """график по нескольким метрикам"""
    metric_ids = block.get("metric_ids", [])
    chart_type = block.get("chart_type", "bar")
 
    labels, values = [], []
    for mid in metric_ids:
        metric_data = metrics.get(mid)
        if metric_data and metric_data.get("value") is not None:
            labels.append(metric_data.get("name", mid))
            values.append(float(metric_data["value"]))
 
    if not labels:
        return
 
    title = ", ".join(labels)
    buf = make_chart(chart_type, labels, values, title)
    doc.add_picture(buf, width=Inches(5.5))
 
 
def render_answers_block(
    doc: DocxDocument,
    block: dict,
    answers_by_question: dict,
    questions_by_id: dict,
    show_all: bool = False,
) -> None:
    """ответы на конкретные вопросы (или все если show_all_answers)"""
    if show_all:
        question_ids = list(questions_by_id.keys())
    else:
        question_ids = [uuid.UUID(qid) for qid in block.get("question_ids", [])]
 
    for qid in question_ids:
        question = questions_by_id.get(qid)
        answer = answers_by_question.get(qid)
        if not question or not answer:
            continue
 
        p = doc.add_paragraph()
        p.add_run(f"{question.text}  ").bold = True
 
        raw = answer.value.get("answer")
        p.add_run(format_answer(raw, question))
 
 
def format_answer(raw, question: Question) -> str:
    """форматирует сырой ответ в нормальную строку"""
    cfg = question.config or {}
 
    if question.type in (QuestionType.SINGLE, QuestionType.YESNO):
        for opt in cfg.get("options", []):
            if opt["id"] == raw:
                return opt.get("text", str(raw))
        return str(raw)
 
    if question.type == QuestionType.MULTIPLE:
        selected_ids = set(raw) if isinstance(raw, list) else set()
        texts = [
            opt["text"] for opt in cfg.get("options", [])
            if opt["id"] in selected_ids
        ]
        return ", ".join(texts) if texts else str(raw)
 
    return str(raw)
 
 
# =====================================
# СБОРКА ДОКУМЕНТА
# =====================================

COLOR_PRIMARY   = RGBColor(0x57, 0xCE, 0x5F)
COLOR_SECONDARY = RGBColor(0x2A, 0x37, 0x40)
COLOR_LIGHT     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_DARK      = RGBColor(0x14, 0x42, 0x10)


def set_paragraph_shading(paragraph, hex_color: str) -> None:
    """закрашивает фон параграфа"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_horizontal_rule(doc) -> None:
    """добавляет горизонтальную линию"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "57CE5F")
    pb.append(bottom)
    pPr.append(pb)


def add_title_block(doc, title: str, subtitle: str) -> None:
    """цветной заголовочный блок"""

    # Верхняя полоса
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(p, "57CE5F")
    run = p.add_run("  ")
    run.font.size = Pt(4)

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_shading(p, "F0FFF0") 
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    # Подзаголовок (имя юзера)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, "F0FFF0") 
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_SECONDARY

    # Нижняя полоса
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    set_paragraph_shading(p, "57CE5F")
    run = p.add_run("  ")
    run.font.size = Pt(4)


def styled_paragraph(doc, text: str, bold=False, size=11, color=None, space_before=4, space_after=4):
    """параграф с кастомным стилем"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def build_docx(
    template_blocks: list[dict],
    session,
    metrics: dict,
    answers_by_question: dict,
    questions_by_id: dict,
    title: str,
) -> io.BytesIO:
    """делает сбрку всего что есть"""
    from docx import Document
    doc = Document()

    #поля страницы
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    #заголовочный блок
    add_title_block(doc, title, f"Клиент: {session.client_name}")

    #блоки из шаблона
    for block in template_blocks:
        block_type = block.get("type")

        if block_type == "text":
            content = block.get("content", "")
            if content:
                styled_paragraph(doc, content, size=11, space_before=6, space_after=6)

        elif block_type == "metric":
            metric_id   = block.get("metric_id", "")
            metric_data = metrics.get(metric_id)
            if not metric_data:
                continue

            add_horizontal_rule(doc)

            name           = metric_data.get("name", metric_id)
            value          = metric_data.get("value")
            interpretation = metric_data.get("interpretation", "")

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)

            r = p.add_run(f"{name}: ")
            r.font.name  = "Calibri"
            r.font.size  = Pt(12)
            r.font.bold  = True
            r.font.color.rgb = COLOR_PRIMARY

            r2 = p.add_run(str(round(value, 2)) if isinstance(value, float) else str(value))
            r2.font.name = "Calibri"
            r2.font.size = Pt(12)
            r2.font.bold = True

            if interpretation:
                pi = doc.add_paragraph()
                pi.paragraph_format.space_before = Pt(0)
                pi.paragraph_format.space_after  = Pt(6)
                ri = pi.add_run(f"  {interpretation}")
                ri.font.name   = "Calibri"
                ri.font.size   = Pt(10)
                ri.font.italic = True
                ri.font.color.rgb = COLOR_SECONDARY

        elif block_type == "chart":
            render_chart_block(doc, block, metrics)
            doc.add_paragraph()

        elif block_type == "answers":
            add_horizontal_rule(doc)
            show_all = block.get("show_all_answers", False)

            if show_all:
                question_ids = list(questions_by_id.keys())
            else:
                question_ids = [uuid.UUID(qid) for qid in block.get("question_ids", [])]

            for qid in question_ids:
                question = questions_by_id.get(qid)
                answer   = answers_by_question.get(qid)
                if not question or not answer:
                    continue

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(1)
                r = p.add_run(question.text)
                r.font.name  = "Calibri"
                r.font.size  = Pt(10)
                r.font.bold  = True
                r.font.color.rgb = COLOR_DARK

                raw = answer.value.get("answer")
                pa  = doc.add_paragraph()
                pa.paragraph_format.space_before = Pt(0)
                pa.paragraph_format.space_after  = Pt(4)
                pa.paragraph_format.left_indent  = Cm(0.5)
                ra  = pa.add_run(format_answer(raw, question))
                ra.font.name = "Calibri"
                ra.font.size = Pt(10)

    #нижний колонтитул с датой
    from datetime import datetime
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run(f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run.font.name  = "Calibri"
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
 
 
# =====================================
# ПУБЛИЧНЫЕ ФУНКЦИИ
# =====================================
 
async def generate_guest_report(db: AsyncSession, session_id: uuid.UUID) -> io.BytesIO:
    """
    генерирует docx отчёт для гостя
    """
    data = await load_session_data(db, session_id)
    session: Session = data["session"]
    test: Test = data["test"]
 
    #проверяем что психолог разрешил показывать отчёт клиенту
    if not test.show_report_to_client:
        raise PermissionError("Психолог не разрешил показывать отчёт клиенту")
 
    template = test.report_template or {}
    client_blocks = template.get("client", [])
 
    return build_docx(
        template_blocks=client_blocks,
        session=session,
        metrics=data["metrics"],
        answers_by_question=data["answers_by_question"],
        questions_by_id=data["questions_by_id"],
        title=f"Результаты теста: {test.title}",
    )
 
 
async def generate_psychologist_report(db: AsyncSession, session_id: uuid.UUID) -> io.BytesIO:
    """
    генерирует docx отчёт для психолога
    """
    data = await load_session_data(db, session_id)
    session: Session = data["session"]
    test: Test = data["test"]
 
    template = test.report_template or {}
    psychologist_blocks = template.get("psychologist", [])
 
    return build_docx(
        template_blocks=psychologist_blocks,
        session=session,
        metrics=data["metrics"],
        answers_by_question=data["answers_by_question"],
        questions_by_id=data["questions_by_id"],
        title=f"Отчёт психолога: {test.title} — {session.client_name}",
    )